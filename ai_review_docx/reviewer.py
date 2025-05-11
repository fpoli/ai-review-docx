from typing import Iterator, Optional
from loguru import logger
import docx
from diskcache import Cache
import litellm
from .comments import add_formatted_comment
from .utils import preview, colored_console_diff, formatted_diff_for_docx


class DocxReviewer:
    """
    Class for reviewing and correcting DOCX documents using an LLM.
    """

    def __init__(self, document_path: str, model_name: str, cache_location: str, api_key: Optional[str] = None,
                 base_url: Optional[str] = None, context: Optional[str] = None):
        """
        Initialize the DocxReviewer.

        Args:
            document_path: Path to the DOCX document to review.
            model_name: The name of the model to be used (e.g., "ollama/gemma:7b").
            cache_location: Directory path where to store the cache of model responses.
            api_key: The API key for the LLM provider.
            base_url: The base URL for the LLM provider API.
            context: Optional context to add to the review prompt.
        """
        logger.info(f"Initializing DocxReviewer on document '{document_path}' using model '{model_name}'")
        self.model_name = model_name
        self.document = docx.Document(document_path)
        self.context = context
        self.cache = Cache(cache_location)
        self.api_key = api_key
        self.base_url = base_url

    def ask_llm(self, prompt: str) -> str:
        """
        Sends a prompt to the LLM and returns the response, using a cache.

        :param prompt: The prompt to send to the language model.
        :return: The response from the language model.
        """
        cache_key = (self.model_name, prompt)

        with self.cache as cache:
            if cache_key in cache:
                return cache[cache_key]
            else:
                response = litellm.completion(
                    model=self.model_name,
                    messages=[{"role": "user", "content": prompt}],
                    api_key=self.api_key,
                    base_url=self.base_url,
                )
                response = response.choices[0].message.content.strip()
                cache[cache_key] = response
                return response

    def report_paragraph_changes(self, paragraph_id: str, paragraph: docx.text.paragraph.Paragraph,
                                 corrected_text: Optional[str]):
        """
        Reports the changes made to a paragraph, adding a comment to the document.
        """
        if corrected_text is None:
            logger.debug(f"Review of {paragraph_id}: No changes proposed by LLM.")
            return

        original_text = paragraph.text.strip()
        if original_text == corrected_text:
            logger.debug(f"Review of {paragraph_id}: No changes proposed by LLM.")
            return

        logger.warning(f"Review of {paragraph_id}: Change proposed.")
        logger.warning(f"  Original:  '{original_text}'")
        logger.warning(f"  Suggested: '{corrected_text}'")
        logger.warning(f"  Diff:      '{colored_console_diff(original_text, corrected_text)}'")

        # Add the colored diff to the docx comment
        formatted_runs = formatted_diff_for_docx(original_text, corrected_text)
        add_formatted_comment(self.document, paragraph, "Reviewer", formatted_runs)


    def review_paragraph(self, paragraph_id: str, paragraph: docx.text.paragraph.Paragraph):
        """
        Reviews text using LLM, returning the corrected text.

        Args:
            paragraph_id: An identifier of the paragraph to be reviewed.
            paragraph: The paragraph to be reviewed.
        """
        paragraph_text = paragraph.text.strip()
        logger.info(f"Reviewing {paragraph_id}: '{preview(paragraph_text)}'")

        try:
            if paragraph_text:
                prompt = (
                    "Review the following text for obvious grammatical errors and spelling mistakes. "
                    "If you find mistakes, return ONLY the corrected text. "
                    "If the text is already correct, return ONLY the original text unchanged."
                )
                if self.context:
                    prompt += f"\n\n{self.context}"
                prompt += f"\n\nText:\n\n{paragraph_text}"
                corrected_text = self.ask_llm(prompt)
            else:
                corrected_text = paragraph_text
            self.report_paragraph_changes(paragraph_id, paragraph, corrected_text)
        except Exception as e:
            logger.error(f"Unexpected error during LLM review: {e}", exc_info=True)

    def review_table(self, table_id: str, table: docx.table.Table):
        """
        Reviews a table using LLM, returning (True, corrected_text) if changed, else (False, None).
        """
        logger.info(f"Reviewing table {table_id}.")

        for cell_index, cell in enumerate(table_itercells(table)):
            for paragraph_index, paragraph in enumerate(cell.paragraphs):
                paragraph_id = f"table {table_id}, cell {cell_index}, paragraph {paragraph_index}/{len(cell.paragraphs)}"
                self.review_paragraph(paragraph_id, paragraph)

    def review(self):
        """
        Reviews the document for grammatical and spelling errors.
        """
        total_paragraphs = len(self.document.paragraphs)
        logger.info(f"Processing {total_paragraphs} paragraphs.")

        for i, paragraph in enumerate(self.document.paragraphs):
            self.review_paragraph(f"paragraph {i}/{total_paragraphs}", paragraph)

        total_tables = len(self.document.tables)
        logger.info(f"Processing {total_tables} tables.")

        for i, table in enumerate(self.document.tables):
            self.review_table(f"{i}/{total_tables}", table)

    def save(self, output_path: str):
        """
        Saves the reviewed document to a new DOCX file.

        Args:
            output_path: Path to the output DOCX file.
        """
        self.document.save(output_path)


def table_itercells(table: docx.table.Table) -> Iterator[docx.table._Cell]:
    """
    Iterate over the cells of a table, without revisiting cells merged vertically or horizontally.

    Source: https://stackoverflow.com/a/78935428/2491528
    """
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            # check if the cell equals the previous cell either horizontally or vertically
            if (r > 0 and c < len(table.rows[r - 1].cells) and cell._tc is table.rows[r - 1].cells[c]._tc) \
                or (c > 0 and cell._tc is row.cells[c - 1]._tc):
                continue
            yield cell
